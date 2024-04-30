from docx import Document
from docx.shared import Inches
from functions import *
import os
from docx.enum.section import WD_ORIENTATION

def create_document(folder_path, output_file):
    document = Document()

    document.add_heading("Prosjeci čišćenja", 0)
    section = document.sections[0]
    section.orientation = WD_ORIENTATION.PORTRAIT
    section.page_width = Inches(8.3)
    section.page_height= Inches(11.7)

    document.add_section()

    section = document.sections[1]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    section.page_width = Inches(11.7)
    section.page_height = Inches(8.3)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.33)
    section.bottom_margin = Inches(1)

    files = os.listdir(folder_path)
    for file_name in files:
        file_path = os.path.join(folder_path, file_name)
        print(file_path)
        if os.path.isfile(file_path):
            document.add_heading(file_name, level = 1)
            document.add_picture(file_path, width = Inches(9.42),height = Inches(7.7))

    document.save(f"{output_file}.docx")
